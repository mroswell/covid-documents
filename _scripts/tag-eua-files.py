import os
import json
import time
import re
from pathlib import Path
from typing import List, Dict, Tuple, Optional
import anthropic
from PIL import Image
import pytesseract
import PyPDF2
from docx import Document
import pandas as pd
from tqdm import tqdm
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from google.auth.transport.requests import Request
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
import io
import pickle
import tempfile

class GoogleDriveFileTagger:
    def __init__(self, claude_api_key: str):
        """Initialize the auto-tagger with API keys"""
        self.claude = anthropic.Anthropic(api_key=claude_api_key)
        self.supported_extensions = {'.pdf', '.jpg', '.jpeg', '.png', '.docx', '.doc', '.xml', '.xpt', '.xsl', '.jmp'}
        self.module_patterns = ['M1', 'M2', 'M3', 'M4', 'M5', 'M6']
        self.temp_dir = tempfile.mkdtemp()
        
        # Google Drive setup
        self.SCOPES = ['https://www.googleapis.com/auth/drive.readonly']
        self.service = self.authenticate_google_drive()
    
    def authenticate_google_drive(self):
        """Authenticate and return Google Drive service"""
        creds = None
        
        # Token file stores the user's access and refresh tokens
        if os.path.exists('token.pickle'):
            with open('token.pickle', 'rb') as token:
                creds = pickle.load(token)
        
        # If there are no (valid) credentials available, let the user log in
        if not creds or not creds.valid:
            if creds and creds.expired and creds.refresh_token:
                creds.refresh(Request())
            else:
                flow = InstalledAppFlow.from_client_secrets_file(
                    'credentials.json', self.SCOPES)
                creds = flow.run_local_server(port=0)
            
            # Save the credentials for the next run
            with open('token.pickle', 'wb') as token:
                pickle.dump(creds, token)
        
        return build('drive', 'v3', credentials=creds)
    
    def get_folder_contents(self, folder_id: str) -> List[Dict]:
        """Get all files in a Google Drive folder"""
        files = []
        page_token = None
        
        while True:
            response = self.service.files().list(
                q=f"'{folder_id}' in parents",
                spaces='drive',
                fields='nextPageToken, files(id, name, mimeType, webViewLink)',
                pageToken=page_token
            ).execute()
            
            files.extend(response.get('files', []))
            page_token = response.get('nextPageToken', None)
            
            if page_token is None:
                break
        
        return files
    
    def download_file_temporarily(self, file_id: str, filename: str, mime_type: str) -> Optional[str]:
        """Download file to temp directory for processing"""
        try:
            # Determine export mime type for Google Docs
            export_mime_type = None
            if mime_type == 'application/vnd.google-apps.document':
                export_mime_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                filename = filename + '.docx' if not filename.endswith('.docx') else filename
            
            # Download file
            if export_mime_type:
                request = self.service.files().export_media(fileId=file_id, mimeType=export_mime_type)
            else:
                request = self.service.files().get_media(fileId=file_id)
            
            temp_path = os.path.join(self.temp_dir, filename)
            fh = io.FileIO(temp_path, 'wb')
            downloader = MediaIoBaseDownload(fh, request)
            
            done = False
            while done is False:
                status, done = downloader.next_chunk()
            
            return temp_path
            
        except Exception as e:
            print(f"Error downloading {filename}: {str(e)}")
            return None
    
    def extract_text_from_pdf(self, file_path: str, max_pages: int = 7) -> Tuple[str, int]:
        """Extract text from PDF (first 7 pages for efficiency) and estimate total pages"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                total_pages = len(pdf_reader.pages)
                text = ""
                pages_to_read = min(total_pages, max_pages)
                
                for page_num in range(pages_to_read):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\n"
                
                # If we have actual page count, use it
                if total_pages > 0:
                    return text[:5000], total_pages
                else:
                    # Fallback to character-based estimation
                    estimated_pages = max(1, len(text) // 3000)
                    return text[:5000], estimated_pages
        except Exception as e:
            return f"Error reading PDF: {str(e)}", 0
    
    def extract_text_from_docx(self, file_path: str, max_paragraphs: int = 50) -> Tuple[str, int]:
        """Extract text from DOCX file (first 50 paragraphs) and estimate pages"""
        try:
            doc = Document(file_path)
            paragraphs = []
            
            for i, para in enumerate(doc.paragraphs):
                if i >= max_paragraphs:
                    break
                if para.text.strip():
                    paragraphs.append(para.text)
            
            full_text = "\n".join(paragraphs)
            
            # Get all text for page estimation
            all_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            # Estimate pages based on characters: ~3000 chars per page
            estimated_pages = max(1, len(all_text) // 3000)
            
            return full_text[:5000], estimated_pages
        except Exception as e:
            return f"Error reading DOCX: {str(e)}", 0
    
    def extract_text_from_image(self, file_path: str) -> Tuple[str, int]:
        """Extract text from image using OCR"""
        try:
            # Basic image metadata
            img = Image.open(file_path)
            metadata = f"Image size: {img.size}, Format: {img.format}\n"
            
            # OCR (optional - requires tesseract installed)
            try:
                text = pytesseract.image_to_string(img)
                content = metadata + f"OCR Text: {text[:2000]}"
            except:
                content = metadata + "No OCR available"
            
            return content, 1  # Images are 1 page
        except Exception as e:
            return f"Error reading image: {str(e)}", 0
    
    def extract_text_from_xml(self, file_path: str, max_chars: int = 5000) -> Tuple[str, int]:
        """Extract text from XML file and estimate pages"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                content = file.read()
                # Simple text extraction from XML
                import re
                text = re.sub('<[^<]+?>', ' ', content)
                # Remove extra whitespace
                text = ' '.join(text.split())
                # Estimate pages based on text length: ~3000 chars per page
                estimated_pages = max(1, len(text) // 3000)
                return text[:max_chars], estimated_pages
        except Exception as e:
            return f"Error reading XML: {str(e)}", 0
    
    def extract_file_content(self, file_path: str) -> Tuple[str, str, int]:
        """Extract content based on file type and return page count"""
        path = Path(file_path)
        extension = path.suffix.lower()
        
        if extension == '.pdf':
            content, pages = self.extract_text_from_pdf(file_path)
        elif extension in ['.docx', '.doc']:
            content, pages = self.extract_text_from_docx(file_path)
        elif extension in ['.jpg', '.jpeg', '.png']:
            content, pages = self.extract_text_from_image(file_path)
        elif extension in ['.xml', '.xpt', '.xsl', '.jmp']:
            content, pages = self.extract_text_from_xml(file_path)
        else:
            content = "Unsupported file type"
            pages = 0
        
        # Add file metadata
        file_info = f"Filename: {path.name}\nEstimated pages: {pages}\n"
        
        return file_info + content, extension, pages
    
    def check_module_tags(self, filename: str) -> List[str]:
        """Check filename for module tags (M1-M6)"""
        module_tags = []
        for module in self.module_patterns:
            if module in filename.upper():
                module_tags.append(module)
        return module_tags
    
    def check_exemption_tag(self, content: str) -> bool:
        """Check if content contains exemption keywords"""
        exemption_keywords = ['exempt', 'exemption']
        content_lower = content.lower()
        return any(keyword in content_lower for keyword in exemption_keywords)
    
    def check_exclusion(self, content: str) -> bool:
        """Check if content contains exclusion keyword"""
        return 'exclusion' in content.lower()
    
    def standardize_tags(self, tags: List[str]) -> List[str]:
        """Standardize tag capitalization according to rules"""
        # Special format terms
        special_formats = {
            'mrna': 'mRNA',
            'modrna': 'modRNA',
            'ecrf': 'eCRF',
            'hepo': 'hEPO',
            'biontech': 'BioNTech',
            'gmbh': 'GmbH',
            'mtb': 'Mtb'
        }
        
        # All caps terms
        all_caps = {
            'fda', 'eua', 'cber', 'cdc', 'who', 'nih', 'niehs', 'cioms', 'vaers', 'bimo',
            'glp', 'gmp', 'cmc', 'ocbq/dmpq/mrbi', 'cber/ovrr/dvp/ldv', 'covid-19', 'covid',
            'sars-cov-2', 'sars-cov', 'mers-cov', 'crf', 'crfs', 'rna', 'dna', 'adsl',
            'adae', 'adva', 'adcevd', 'cdisc', 'sas', 'sdtm', 'rt-pcr', 'pcr', 'naat',
            'bmi', 'ecg', 'thc', 'hek293t', 'balb/c', 'gs1', 'ndc', 'bnt162b2', 'bnt162b1',
            'bnt162', 'ind', 'pf-07302048', 'c4591001', 'm1', 'm2', 'm3', 'm4', 'm5', 'm6',
            'alc-0315', 'alc-0159', 'dspc', 'ace2', 'dpp4', 'rbd', 'vrbpac', 'vaed', 'mis',
            'evali', 'rsv', 'fi-rsv', 'activ', 'barda', 'dart', 'gmfr', 'aai', 'isaric',
            'cepi', 'mcdc', 'sbu', 'foia', 'lnp', 'peg', 'tmprss2', 'icos', 'nepa', 'ceq',
            'ovrr', 'ib', 'pvp', 'qc', 'us', 'usa', 'uk', 'eu/eea', 'ecdc', 'rs', 'cy',
            'lpt', 'suny', 'va'
        }
        
        standardized = []
        for tag in tags:
            # Handle special formats first
            tag_lower = tag.lower()
            if tag_lower in special_formats:
                standardized.append(special_formats[tag_lower])
            elif tag_lower in all_caps:
                standardized.append(tag_lower.upper())
            elif tag_lower == 'slipsheet':
                standardized.append('Slipsheet')
            else:
                # Handle phrases with special terms
                words = tag.split()
                new_words = []
                for word in words:
                    word_lower = word.lower()
                    if word_lower in special_formats:
                        new_words.append(special_formats[word_lower])
                    elif word_lower in all_caps:
                        new_words.append(word_lower.upper())
                    else:
                        # Proper case
                        new_words.append(word.capitalize())
                standardized.append(' '.join(new_words))
        
        return standardized
    
    def generate_tags_and_title_with_claude(self, file_content: str, filename: str, file_type: str) -> Tuple[List[str], str, str, str, List[str], bool]:
        """Use Claude API to generate relevant tags, guess document title, date, document type, people mentioned, and password protection"""
        
        # Adjust tag count for JPG files
        tag_count = "5-8" if file_type in ['.jpg', '.jpeg', '.png'] else "exactly 13"
        
        prompt = f"""Analyze this file content and provide:
1. A best guess for the document's title (what this document would be called)
2. A best guess for the document's date (when it was created or pertains to)
3. The document type/purpose (e.g., report, presentation, invoice, memo, form, etc.) - just one or two words
4. List of people mentioned by name in the document
5. Whether this file appears to be password-protected or encrypted (based on content patterns)
6. Generate {tag_count} relevant tags for categorization (DO NOT include document type in tags)

Filename: {filename}
File type: {file_type}
Content excerpt:
{file_content[:2000]}

For the title: Look for document headers, title pages, or infer from content. Keep it concise (3-7 words).

For the date: Look for dates in the content, headers, or filename. Format as YYYY-MM-DD if full date available, or YYYY-MM, or just YYYY if that's all you can determine. If no date found, return "undated".

For document type: Identify the primary document type (report, memo, presentation, email, form, invoice, etc.)

For people: List any person names mentioned (first name, last name, or full names). Return empty list if none found.

For password protection: Look for patterns like repeated characters, encoding artifacts, or messages about encryption. Return true if likely protected, false otherwise.

For tags (excluding document type), include:
- Main topics or subjects
- Key entities, dates, or categories mentioned
- Department or area if apparent
- Any compliance or regulatory aspects
- Relevant time periods (Q1, 2024, etc.)
- Key organizations mentioned

Return ONLY a JSON object. Example:
{{"title": "Q3 2024 Financial Report", "date": "2024-09", "document_type": "report", "people": ["John Smith", "Jane Doe"], "password_protected": false, "tags": ["financial", "Q3 2024", "sales data", "revenue analysis", "quarterly review", "executive summary", "financial planning", "budget", "third quarter", "earnings", "corporate finance", "fiscal review", "performance metrics"]}}"""

        try:
            response = self.claude.messages.create(
                model="claude-3-haiku-20240307",  # Using Haiku for cost efficiency
                max_tokens=300,
                temperature=0.3,
                messages=[{"role": "user", "content": prompt}]
            )
            
            # Parse the response
            result_json = response.content[0].text.strip()
            
            # Try to fix common JSON issues
            # Remove any text before the first {
            if '{' in result_json:
                result_json = result_json[result_json.index('{'):]
            # Remove any text after the last }
            if '}' in result_json:
                result_json = result_json[:result_json.rindex('}')+1]
            
            # Fix common JSON formatting issues
            result_json = result_json.replace('\n', ' ')  # Remove newlines
            result_json = result_json.replace('\\', '\\\\')  # Escape backslashes
            
            try:
                result = json.loads(result_json)
            except json.JSONDecodeError as e:
                print(f"JSON parsing error for {filename}: {str(e)}")
                print(f"Raw response: {result_json[:200]}...")
                # Provide default values
                result = {
                    'title': filename.replace('_', ' ').replace('.pdf', '').replace('.docx', ''),
                    'date': 'undated',
                    'document_type': 'document',
                    'people': [],
                    'password_protected': False,
                    'tags': self.generate_fallback_tags(filename, file_content)
                }
            
            title = result.get('title', 'Untitled Document')
            date_guess = result.get('date', 'undated')
            document_type = result.get('document_type', 'unknown')
            people = result.get('people', [])
            password_protected = result.get('password_protected', False)
            tags = result.get('tags', [])
            
            # Ensure minimum tags for non-image files
            if file_type not in ['.jpg', '.jpeg', '.png'] and len(tags) < 13:
                tags.extend(['needs-review'] * (13 - len(tags)))
            
            return tags, title, date_guess, document_type, people, password_protected
            
        except Exception as e:
            print(f"Error generating tags for {filename}: {str(e)}")
            # Generate fallback tags based on filename
            fallback_tags = self.generate_fallback_tags(filename, file_content)
            return fallback_tags, filename.replace('_', ' ').replace('.pdf', ''), "undated", "unknown", [], False
    
    def generate_fallback_tags(self, filename: str, content: str) -> List[str]:
        """Generate basic tags when Claude API fails"""
        tags = []
        
        # Extract from filename
        filename_lower = filename.lower()
        if 'report' in filename_lower:
            tags.append('report')
        if any(year in filename_lower for year in ['2020', '2021', '2022', '2023', '2024', '2025']):
            for year in ['2020', '2021', '2022', '2023', '2024', '2025']:
                if year in filename_lower:
                    tags.append(year)
        
        # Check content for common terms
        content_lower = content.lower()
        if 'clinical' in content_lower:
            tags.append('clinical')
        if 'trial' in content_lower:
            tags.append('trial')
        if 'covid' in content_lower or 'sars' in content_lower:
            tags.append('COVID-19')
        if 'vaccine' in content_lower:
            tags.append('vaccine')
        if 'fda' in content_lower:
            tags.append('FDA')
        if 'study' in content_lower:
            tags.append('study')
        if 'research' in content_lower:
            tags.append('research')
        if 'data' in content_lower:
            tags.append('data')
        
        # Pad to minimum
        while len(tags) < 13:
            tags.append('review-needed')
        
        return tags[:13]
    
    def process_google_drive_folder(self, folder_id: str, folder_name: str, display_name: str) -> List[Dict]:
        """Process all files in a Google Drive folder"""
        results = []
        
        # Get files from Google Drive
        print(f"\nFetching files from {folder_name}...")
        drive_files = self.get_folder_contents(folder_id)
        
        # Filter for supported file types
        supported_files = []
        for file in drive_files:
            filename = file['name']
            if any(filename.lower().endswith(ext) for ext in self.supported_extensions):
                supported_files.append(file)
            # Also check Google Docs
            elif file['mimeType'] == 'application/vnd.google-apps.document':
                supported_files.append(file)
            # Also check for .doc files
            elif filename.lower().endswith('.doc'):
                supported_files.append(file)
        
        print(f"Processing {len(supported_files)} supported files in {folder_name}...")
        
        # TEST MODE: Process only first 7 files
        # supported_files = supported_files[:7]
        # print(f"TEST MODE: Processing only {len(supported_files)} files")
        
        for file in tqdm(supported_files):
            try:
                filename = file['name']
                file_id = file['id']
                google_drive_link = file['webViewLink']
                
                # Download file temporarily
                temp_path = self.download_file_temporarily(file_id, filename, file['mimeType'])
                
                if temp_path:
                    # Extract content
                    content, file_type, page_count = self.extract_file_content(temp_path)
                    
                    # Generate AI tags, title, date, and metadata
                    ai_tags, document_title, document_date, document_type, people_mentioned, password_protected = \
                        self.generate_tags_and_title_with_claude(content, filename, file_type)
                    
                    # Standardize AI tags
                    ai_tags = self.standardize_tags(ai_tags)
                    
                    # Check for module tags in filename
                    module_tags = self.check_module_tags(filename)
                    
                    # Check for exemption in content
                    has_exemption = self.check_exemption_tag(content)
                    
                    # Check for exclusion in content
                    has_exclusion = self.check_exclusion(content)
                    
                    # Build exemption/exclusion tags
                    special_tags = []
                    if has_exemption:
                        special_tags.append("Exemption")
                    if has_exclusion:
                        special_tags.append("Exclusion")
                    
                    # Add module tags to main tags
                    module_tag_list = [f"M{num}" for num in module_tags]
                    
                    # Combine all tags
                    all_tags = list(set(ai_tags + special_tags + module_tag_list))
                    
                    # Format people mentioned
                    people_mentioned_str = ', '.join(people_mentioned) if people_mentioned else ''
                    
                    # Store result
                    result = {
                        'filename': filename,
                        'title': document_title,
                        'date': document_date,
                        'google_drive_link': google_drive_link,
                        'folder': display_name,  # Use the shortened display name
                        'file_type': file_type.upper().replace('.', '') if file_type else 'UNKNOWN',
                        'page_count': page_count,
                        'module': module_tags[0] if module_tags else '',  # First module found or empty
                        'document_type': document_type,
                        'people_mentioned': people_mentioned_str,
                        'tags': all_tags,
                        'has_exemption': has_exemption,
                        'has_exclusion': has_exclusion,
                        'password_protected': password_protected,
                        'processed': True
                    }
                    results.append(result)
                    
                    # Clean up temp file
                    os.remove(temp_path)
                    
                    # Rate limiting for API
                    time.sleep(0.5)  # Adjust based on your API limits
                else:
                    raise Exception("Failed to download file")
                    
            except Exception as e:
                print(f"Error processing {file['name']}: {str(e)}")
                results.append({
                    'filename': file['name'],
                    'title': 'Unknown Document',
                    'date': 'undated',
                    'google_drive_link': file.get('webViewLink', ''),
                    'folder': display_name,  # Use the shortened display name
                    'file_type': 'ERROR',
                    'page_count': 0,
                    'module': '',  # Empty for error cases
                    'document_type': 'unknown',
                    'people_mentioned': '',
                    'tags': ['error', 'unprocessed'],
                    'has_exemption': False,
                    'has_exclusion': False,
                    'password_protected': False,
                    'processed': False
                })
        
        return results
    
    def save_to_csv(self, results_dict: Dict[str, List[Dict]], output_file: str = "google_drive_tagged_files.csv"):
        """Save results to CSV with separate tabs for each folder"""
        with pd.ExcelWriter(output_file.replace('.csv', '.xlsx'), engine='openpyxl') as writer:
            for folder_name, results in results_dict.items():
                if results:
                    df = pd.DataFrame(results)
                    # Convert tags list to comma-separated string
                    df['tags'] = df['tags'].apply(lambda x: ', '.join(x))
                    # Write to a separate sheet
                    sheet_name = folder_name[:31]  # Excel sheet name limit
                    df.to_excel(writer, sheet_name=sheet_name, index=False)
        
        print(f"\nResults saved to {output_file.replace('.csv', '.xlsx')}")
        
        # Also save as CSV for easy Airtable import
        all_results = []
        for folder_name, results in results_dict.items():
            all_results.extend(results)
        
        df_all = pd.DataFrame(all_results)
        df_all['tags'] = df_all['tags'].apply(lambda x: ', '.join(x) if isinstance(x, list) else x)
        df_all.to_csv(output_file, index=False)
        print(f"Combined CSV saved to {output_file}")
    
    def generate_report(self, all_results: Dict[str, List[Dict]], output_file: str = "file_tagging_report.txt"):
        """Generate a detailed report of the tagging process"""
        with open(output_file, 'w') as f:
            f.write("Google Drive File Tagging Report\n")
            f.write("=" * 50 + "\n\n")
            f.write(f"Report generated on: {pd.Timestamp.now()}\n\n")
            
            # Overall statistics
            total_files = sum(len(results) for results in all_results.values())
            successful = sum(sum(1 for r in results if r['processed']) for results in all_results.values())
            errors = total_files - successful
            
            f.write("OVERALL STATISTICS\n")
            f.write("-" * 30 + "\n")
            f.write(f"Total files processed: {total_files}\n")
            f.write(f"Successfully tagged: {successful}\n")
            f.write(f"Errors: {errors}\n\n")
            
            # Statistics by file type
            f.write("FILES BY TYPE\n")
            f.write("-" * 30 + "\n")
            file_type_counts = {}
            for results in all_results.values():
                for r in results:
                    ft = r['file_type']
                    file_type_counts[ft] = file_type_counts.get(ft, 0) + 1
            
            for ft, count in sorted(file_type_counts.items()):
                f.write(f"{ft}: {count}\n")
            f.write("\n")
            
            # Per folder statistics
            f.write("FOLDER STATISTICS\n")
            f.write("-" * 30 + "\n")
            for folder_name, results in all_results.items():
                f.write(f"\n{folder_name}:\n")
                f.write(f"  Total files: {len(results)}\n")
                f.write(f"  Successfully tagged: {sum(1 for r in results if r['processed'])}\n")
                f.write(f"  Files with module tags: {sum(1 for r in results if r.get('module'))}\n")
                f.write(f"  Files with exemption: {sum(1 for r in results if r['has_exemption'])}\n")
                f.write(f"  Files with exclusion: {sum(1 for r in results if r.get('has_exclusion', False))}\n")
                f.write(f"  Password protected: {sum(1 for r in results if r.get('password_protected', False))}\n")
                f.write(f"  Errors: {sum(1 for r in results if not r['processed'])}\n")
            
            # Files with exemption
            f.write("\n\nFILES CONTAINING EXEMPTION\n")
            f.write("-" * 30 + "\n")
            exemption_files = []
            for folder_name, results in all_results.items():
                for r in results:
                    if r['has_exemption']:
                        exemption_files.append(f"{r['filename']} ({folder_name})")
            
            if exemption_files:
                for file in sorted(exemption_files):
                    f.write(f"  {file}\n")
            else:
                f.write("  None found\n")
            
            # Password protected files
            f.write("\n\nPOTENTIALLY PASSWORD PROTECTED FILES\n")
            f.write("-" * 30 + "\n")
            protected_files = []
            for folder_name, results in all_results.items():
                for r in results:
                    if r.get('password_protected', False):
                        protected_files.append(f"{r['filename']} ({folder_name})")
            
            if protected_files:
                for file in sorted(protected_files):
                    f.write(f"  {file}\n")
            else:
                f.write("  None detected\n")
            
            # Average tags per file
            f.write("\n\nTAG STATISTICS\n")
            f.write("-" * 30 + "\n")
            all_tag_counts = []
            for results in all_results.values():
                for r in results:
                    if r['processed'] and isinstance(r['tags'], list):
                        all_tag_counts.append(len(r['tags']))
            
            if all_tag_counts:
                avg_tags = sum(all_tag_counts) / len(all_tag_counts)
                f.write(f"Average tags per file: {avg_tags:.1f}\n")
                f.write(f"Min tags: {min(all_tag_counts)}\n")
                f.write(f"Max tags: {max(all_tag_counts)}\n")
            
            f.write("\n" + "=" * 50 + "\n")
            f.write("Report generation complete.\n")
        
        print(f"\nDetailed report saved to: {output_file}")
    
    def cleanup(self):
        """Clean up temporary directory"""
        import shutil
        shutil.rmtree(self.temp_dir)

# Example usage
if __name__ == "__main__":
    # Initialize the tagger
    tagger = GoogleDriveFileTagger(
        claude_api_key="" // Put your Claude API key here. 
    )
    
    # Define your Google Drive folders
    folders = {
        "pd-eua-production-051925": {
            "id": "1T2i_mRlujFpozqcqgmGEh5V-iCTRT0YL",
            "display_name": "eua-051925"
        },
        "pd-eua-production-063025": {
            "id": "140H5GgyNdionOk9M6uJfo0YwESRubOXq", 
            "display_name": "eua-063025"
        }
    }
    
    # Process each folder
    all_results = {}
    
    for folder_name, folder_info in folders.items():
        results = tagger.process_google_drive_folder(
            folder_info["id"], 
            folder_name, 
            folder_info["display_name"]
        )
        all_results[folder_info["display_name"]] = results
    
    # Save results
    tagger.save_to_csv(all_results, "google_drive_tagged_files.csv")
    
    # Generate detailed report
    tagger.generate_report(all_results)
    
    # Cleanup
    tagger.cleanup()
    
    # Print summary
    print(f"\nProcessing complete!")
    for folder_name, results in all_results.items():
        print(f"\n{folder_name}:")
        print(f"  Total files: {len(results)}")
        print(f"  Successfully tagged: {sum(1 for r in results if r['processed'])}")
        print(f"  Files with module tags: {sum(1 for r in results if r.get('module'))}")
        print(f"  Files with exemption: {sum(1 for r in results if r['has_exemption'])}")
        print(f"  Errors: {sum(1 for r in results if not r['processed'])}")
